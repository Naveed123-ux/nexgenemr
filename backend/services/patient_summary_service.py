from sqlalchemy.orm import Session, joinedload
from models.patient_summary_model import PatientSummary
from models.user_model import User
from models.patient_profile_model import PatientProfile
from models.appointment_model import Appointment
from models.soap_note_model import SoapNote
from models.prescription_model import Prescription
from models.lab_result_model import LabResult
from models.clinical_data_model import Vitals, MedicalHistory
from schemas.patient_summary_schema import PatientSummaryCreate, PatientSummaryUpdate
from services.gemini_service import GeminiService
from fastapi import HTTPException
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import re


class PatientSummaryService:
    
    @staticmethod
    def clean_markdown(text: str) -> str:
        """Remove markdown formatting from text for PDF/Word display"""
        if not text:
            return text
        
        # Remove bold markers (**text** or __text__)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        
        # Remove italic markers (*text* or _text_)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        
        # Remove code markers (`text`)
        text = re.sub(r'`(.+?)`', r'\1', text)
        
        # Remove headers (# text)
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
        
        return text
    
    @staticmethod
    def get_patient_context_for_summary(db: Session, patient_user_id: int) -> dict:
        """Gather all patient information for summary generation"""
        
        # Get patient user
        patient_user = db.query(User).filter(User.id == patient_user_id).first()
        if not patient_user:
            raise HTTPException(status_code=404, detail="Patient user not found")
        
        # Get patient profile
        patient_profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == patient_user_id
        ).first()
        
        if not patient_profile:
            raise HTTPException(status_code=404, detail="Patient profile not found")
        
        # Get appointments with eager loading
        appointments = db.query(Appointment).options(
            joinedload(Appointment.slot),
            joinedload(Appointment.doctor),
            joinedload(Appointment.soap_note),
            joinedload(Appointment.vitals)
        ).filter(
            Appointment.patient_profile_id == patient_profile.id
        ).order_by(Appointment.id.desc()).limit(10).all()
        
        # Get prescriptions
        prescriptions = db.query(Prescription).filter(
            Prescription.patient_user_id == patient_user_id
        ).all()
        
        # Get lab results
        lab_results = db.query(LabResult).filter(
            LabResult.patient_user_id == patient_user_id
        ).order_by(LabResult.id.desc()).limit(5).all()
        
        # Get medical history
        medical_history = db.query(MedicalHistory).filter(
            MedicalHistory.patient_profile_id == patient_profile.id
        ).first()
        
        # Convert medical history to dict for JSON serialization
        medical_history_dict = None
        if medical_history:
            medical_history_dict = {
                'allergies': medical_history.allergies if medical_history.allergies else [],
                'current_medications': medical_history.current_medications if medical_history.current_medications else [],
                'past_medical_history': medical_history.past_medical_history if medical_history.past_medical_history else []
            }
        
        # Get latest vitals
        latest_vitals = None
        for apt in appointments:
            if apt.vitals:
                latest_vitals = apt.vitals
                break
        
        # Calculate patient age
        patient_age = "Unknown"
        if patient_profile.subscriber_dob:
            from datetime import date
            today = date.today()
            dob = patient_profile.subscriber_dob
            patient_age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
        
        return {
            'patient_user': patient_user,
            'patient_profile': patient_profile,
            'patient_name': f"{patient_user.first_name} {patient_user.last_name}",
            'patient_age': patient_age,
            'appointments': appointments,
            'prescriptions': prescriptions,
            'lab_results': lab_results,
            'medical_history': medical_history_dict,
            'latest_vitals': latest_vitals
        }
    
    @staticmethod
    def create_patient_summary(
        db: Session,
        summary_data: PatientSummaryCreate,
        current_user: User
    ):
        """Create a new AI-generated patient-friendly summary"""
        
        # Verify doctor role
        if current_user.role.name != "Doctor":
            raise HTTPException(
                status_code=403,
                detail="Only doctors can create patient summaries"
            )
        
        # Get patient context
        patient_context = PatientSummaryService.get_patient_context_for_summary(
            db, summary_data.patient_user_id
        )
        
        # Get hospital ID
        hospital_id = None
        if current_user.doctor_profile:
            hospital_id = current_user.doctor_profile.department.hospital_id
        else:
            raise HTTPException(status_code=400, detail="Doctor must be associated with a hospital")
        
        # Generate title if not provided
        title = summary_data.title
        if not title:
            title = f"Visit Summary - {datetime.now().strftime('%B %d, %Y')}"
        
        # Generate AI content using Gemini
        try:
            gemini_service = GeminiService()
            ai_content = gemini_service.generate_patient_friendly_summary(patient_context)
        except Exception as e:
            # If Gemini fails, create basic summary
            ai_content = {
                'ai_generated_summary': f"Summary for {patient_context['patient_name']}. AI generation unavailable: {str(e)}",
                'what_we_found': "Please contact your doctor for information about your visit.",
                'what_it_means': "Your doctor will explain your results.",
                'your_diagnosis': "Please speak with your doctor.",
                'your_treatment_plan': "Your doctor will discuss your treatment plan.",
                'your_medications': f"{len(patient_context['prescriptions'])} medications prescribed.",
                'what_to_watch_for': "Contact your doctor if you have concerns.",
                'next_steps': "Follow up as scheduled.",
                'lifestyle_tips': "Your doctor will provide recommendations.",
                'questions_to_ask': "Write down questions for your next appointment."
            }
        
        # Create patient summary
        patient_summary = PatientSummary(
            patient_user_id=summary_data.patient_user_id,
            doctor_user_id=current_user.id,
            hospital_id=hospital_id,
            title=title,
            summary_date=datetime.utcnow(),
            ai_generated_summary=ai_content.get('ai_generated_summary'),
            what_we_found=ai_content.get('what_we_found'),
            what_it_means=ai_content.get('what_it_means'),
            your_diagnosis=ai_content.get('your_diagnosis'),
            your_treatment_plan=ai_content.get('your_treatment_plan'),
            your_medications=ai_content.get('your_medications'),
            what_to_watch_for=ai_content.get('what_to_watch_for'),
            next_steps=ai_content.get('next_steps'),
            lifestyle_tips=ai_content.get('lifestyle_tips'),
            questions_to_ask=ai_content.get('questions_to_ask'),
            doctor_notes=summary_data.doctor_notes,
            special_instructions=summary_data.special_instructions
        )
        
        db.add(patient_summary)
        db.commit()
        db.refresh(patient_summary)
        
        # Generate PDF and Word documents
        PatientSummaryService.generate_documents(db, patient_summary.id)
        
        return patient_summary
    
    @staticmethod
    def generate_documents(db: Session, summary_id: int):
        """Generate both PDF and Word documents for patient summary"""
        
        summary = db.query(PatientSummary).filter(
            PatientSummary.id == summary_id
        ).first()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Patient summary not found")
        
        # Get patient info
        patient = db.query(User).filter(User.id == summary.patient_user_id).first()
        doctor = db.query(User).filter(User.id == summary.doctor_user_id).first()
        
        # Create output directory
        output_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'patient_summaries')
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate filenames
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"patient_summary_{summary.patient_user_id}_{timestamp}.pdf"
        word_filename = f"patient_summary_{summary.patient_user_id}_{timestamp}.docx"
        pdf_path = os.path.join(output_dir, pdf_filename)
        word_path = os.path.join(output_dir, word_filename)
        
        # Generate PDF
        PatientSummaryService._generate_pdf_document(summary, patient, doctor, pdf_path)
        
        # Generate Word document
        PatientSummaryService._generate_word_document(summary, patient, doctor, word_path)
        
        # Update summary with file paths
        summary.pdf_file_path = f"/static/patient_summaries/{pdf_filename}"
        summary.word_file_path = f"/static/patient_summaries/{word_filename}"
        db.commit()
        
        return summary
    
    @staticmethod
    def _generate_pdf_document(summary: PatientSummary, patient: User, doctor: User, output_path: str):
        """Generate PDF document for patient summary"""
        
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=1*inch, bottomMargin=0.75*inch)
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#1a5490'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1a5490'),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=14
        )
        
        # Add logo if exists
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'static', 'images', 'logos', 'nexgen.png')
        if os.path.exists(logo_path):
            # Maintain aspect ratio - let height auto-adjust
            logo = Image(logo_path, width=2.5*inch)
            story.append(logo)
            story.append(Spacer(1, 0.2*inch))
        
        # Header
        story.append(Paragraph(summary.title, title_style))
        story.append(Paragraph("Your Health Summary in Simple Terms", normal_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Patient information table
        patient_info_data = [
            ['Patient Name:', f"{patient.first_name} {patient.last_name}"],
            ['Summary Date:', summary.summary_date.strftime('%B %d, %Y')],
            ['Your Doctor:', f"Dr. {doctor.first_name} {doctor.last_name}"]
        ]
        
        patient_info_table = Table(patient_info_data, colWidths=[2*inch, 4.5*inch])
        patient_info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e3f2fd')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ]))
        
        story.append(patient_info_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Add sections
        sections = [
            ('What We Found', summary.what_we_found),
            ('What It Means for You', summary.what_it_means),
            ('Your Diagnosis', summary.your_diagnosis),
            ('Your Treatment Plan', summary.your_treatment_plan),
            ('Your Medications', summary.your_medications),
            ('What to Watch For', summary.what_to_watch_for),
            ('Your Next Steps', summary.next_steps),
            ('Lifestyle Tips', summary.lifestyle_tips),
            ('Common Questions', summary.questions_to_ask),
            ('Doctor\'s Additional Notes', summary.doctor_notes),
            ('Special Instructions', summary.special_instructions)
        ]
        
        for section_title, section_content in sections:
            if section_content:
                story.append(Paragraph(section_title, heading_style))
                # Clean markdown formatting
                cleaned_content = PatientSummaryService.clean_markdown(section_content)
                story.append(Paragraph(cleaned_content, normal_style))
                story.append(Spacer(1, 0.15*inch))
        
        # Footer message
        story.append(Spacer(1, 0.5*inch))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph("If you have any questions about this summary, please contact your doctor.", footer_style))
        
        # Add e-signatures section if any signatures exist
        from models.signature_model import UserSignature
        
        if summary.signed_by_doctor_id or summary.signed_by_staff_id or summary.signed_by_admin_id:
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph("Electronic Signatures:", heading_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Doctor signature
            if summary.signed_by_doctor_id and summary.signed_by_doctor:
                doctor_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    doctor_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_doctor_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if doctor_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        doctor_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            sig_img = Image(sig_path, width=2*inch, height=0.75*inch)
                            story.append(sig_img)
                        except:
                            pass
                        
                        doctor_name = f"Dr. {summary.signed_by_doctor.first_name} {summary.signed_by_doctor.last_name}"
                        story.append(Paragraph(f"<b>{doctor_name}</b>", normal_style))
                        if summary.doctor_signed_at:
                            story.append(Paragraph(
                                f"Signed: {summary.doctor_signed_at.strftime('%B %d, %Y at %I:%M %p')}", 
                                normal_style
                            ))
                        story.append(Spacer(1, 0.3*inch))
            
            # Staff signature
            if summary.signed_by_staff_id and summary.signed_by_staff:
                staff_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    staff_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_staff_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if staff_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        staff_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            sig_img = Image(sig_path, width=2*inch, height=0.75*inch)
                            story.append(sig_img)
                        except:
                            pass
                        
                        staff_name = f"{summary.signed_by_staff.first_name} {summary.signed_by_staff.last_name}"
                        story.append(Paragraph(f"<b>{staff_name} (Staff)</b>", normal_style))
                        if summary.staff_signed_at:
                            story.append(Paragraph(
                                f"Signed: {summary.staff_signed_at.strftime('%B %d, %Y at %I:%M %p')}", 
                                normal_style
                            ))
                        story.append(Spacer(1, 0.3*inch))
            
            # Admin signature
            if summary.signed_by_admin_id and summary.signed_by_admin:
                admin_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    admin_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_admin_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if admin_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        admin_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            sig_img = Image(sig_path, width=2*inch, height=0.75*inch)
                            story.append(sig_img)
                        except:
                            pass
                        
                        admin_name = f"{summary.signed_by_admin.first_name} {summary.signed_by_admin.last_name}"
                        story.append(Paragraph(f"<b>{admin_name} (Hospital Admin)</b>", normal_style))
                        if summary.admin_signed_at:
                            story.append(Paragraph(
                                f"Signed: {summary.admin_signed_at.strftime('%B %d, %Y at %I:%M %p')}", 
                                normal_style
                            ))
                        story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(story)
    
    @staticmethod
    def _generate_word_document(summary: PatientSummary, patient: User, doctor: User, output_path: str):
        """Generate Word document for patient summary"""
        from docx.shared import Inches
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add logo if exists
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'static', 'images', 'logos', 'nexgen.png')
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(2.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title = doc.add_heading(summary.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(26, 84, 144)
        
        # Subtitle
        subtitle = doc.add_paragraph("Your Health Summary in Simple Terms")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
        
        # Patient information
        info_para = doc.add_paragraph()
        info_para.add_run(f"Patient Name: ").bold = True
        info_para.add_run(f"{patient.first_name} {patient.last_name}\n")
        info_para.add_run(f"Summary Date: ").bold = True
        info_para.add_run(f"{summary.summary_date.strftime('%B %d, %Y')}\n")
        info_para.add_run(f"Your Doctor: ").bold = True
        info_para.add_run(f"Dr. {doctor.first_name} {doctor.last_name}")
        
        doc.add_paragraph()
        
        # Add sections
        sections_data = [
            ('What We Found', summary.what_we_found),
            ('What It Means for You', summary.what_it_means),
            ('Your Diagnosis', summary.your_diagnosis),
            ('Your Treatment Plan', summary.your_treatment_plan),
            ('Your Medications', summary.your_medications),
            ('What to Watch For', summary.what_to_watch_for),
            ('Your Next Steps', summary.next_steps),
            ('Lifestyle Tips', summary.lifestyle_tips),
            ('Common Questions', summary.questions_to_ask),
            ('Doctor\'s Additional Notes', summary.doctor_notes),
            ('Special Instructions', summary.special_instructions)
        ]
        
        for section_title, section_content in sections_data:
            if section_content:
                heading = doc.add_heading(section_title, level=1)
                heading_run = heading.runs[0]
                heading_run.font.color.rgb = RGBColor(26, 84, 144)
                
                # Clean markdown formatting
                cleaned_content = PatientSummaryService.clean_markdown(section_content)
                content_para = doc.add_paragraph(cleaned_content)
                content_para_format = content_para.paragraph_format
                content_para_format.space_after = Pt(12)
        
        # Footer message
        doc.add_paragraph()
        footer = doc.add_paragraph("If you have any questions about this summary, please contact your doctor.")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        
        # Add e-signatures section if any signatures exist
        from models.signature_model import UserSignature
        
        if summary.signed_by_doctor_id or summary.signed_by_staff_id or summary.signed_by_admin_id:
            doc.add_paragraph()
            doc.add_heading('Electronic Signatures', level=2)
            
            # Doctor signature
            if summary.signed_by_doctor_id and summary.signed_by_doctor:
                doctor_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    doctor_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_doctor_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if doctor_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        doctor_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            doc.add_picture(sig_path, width=Inches(2))
                        except:
                            pass
                        
                        doctor_name = f"Dr. {summary.signed_by_doctor.first_name} {summary.signed_by_doctor.last_name}"
                        p = doc.add_paragraph()
                        p.add_run(doctor_name).bold = True
                        if summary.doctor_signed_at:
                            doc.add_paragraph(f"Signed: {summary.doctor_signed_at.strftime('%B %d, %Y at %I:%M %p')}")
                        doc.add_paragraph()
            
            # Staff signature
            if summary.signed_by_staff_id and summary.signed_by_staff:
                staff_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    staff_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_staff_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if staff_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        staff_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            doc.add_picture(sig_path, width=Inches(2))
                        except:
                            pass
                        
                        staff_name = f"{summary.signed_by_staff.first_name} {summary.signed_by_staff.last_name}"
                        p = doc.add_paragraph()
                        p.add_run(f"{staff_name} (Staff)").bold = True
                        if summary.staff_signed_at:
                            doc.add_paragraph(f"Signed: {summary.staff_signed_at.strftime('%B %d, %Y at %I:%M %p')}")
                        doc.add_paragraph()
            
            # Admin signature
            if summary.signed_by_admin_id and summary.signed_by_admin:
                admin_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    admin_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_admin_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if admin_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        admin_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            doc.add_picture(sig_path, width=Inches(2))
                        except:
                            pass
                        
                        admin_name = f"{summary.signed_by_admin.first_name} {summary.signed_by_admin.last_name}"
                        p = doc.add_paragraph()
                        p.add_run(f"{admin_name} (Hospital Admin)").bold = True
                        if summary.admin_signed_at:
                            doc.add_paragraph(f"Signed: {summary.admin_signed_at.strftime('%B %d, %Y at %I:%M %p')}")
                        doc.add_paragraph()
        
        # Save document
        doc.save(output_path)
    
    @staticmethod
    def get_patient_summary(db: Session, summary_id: int, current_user: User):
        """Get a patient summary by ID"""
        
        summary = db.query(PatientSummary).filter(
            PatientSummary.id == summary_id
        ).first()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Patient summary not found")
        
        # Check permissions
        if current_user.role.name == "Patient":
            # Patient can only view their own summaries
            if summary.patient_user_id != current_user.id:
                raise HTTPException(status_code=403, detail="You can only view your own summaries")
            
            # Mark as viewed
            if not summary.is_viewed_by_patient:
                summary.is_viewed_by_patient = True
                summary.viewed_at = datetime.utcnow()
                db.commit()
        elif current_user.role.name == "Doctor":
            # Doctor can view summaries they created or for their patients
            if summary.doctor_user_id != current_user.id:
                # Check if patient is assigned to this doctor
                patient_profile = db.query(PatientProfile).filter(
                    PatientProfile.user_id == summary.patient_user_id
                ).first()
                if not patient_profile or patient_profile.assigned_doctor_id != current_user.id:
                    raise HTTPException(status_code=403, detail="You can only view summaries for your patients")
        else:
            raise HTTPException(status_code=403, detail="Only patients and doctors can view summaries")
        
        return summary
    
    @staticmethod
    def get_patient_summaries(db: Session, patient_user_id: int, current_user: User):
        """Get all summaries for a patient"""
        
        # Check permissions
        if current_user.role.name == "Patient":
            if patient_user_id != current_user.id:
                raise HTTPException(status_code=403, detail="You can only view your own summaries")
        elif current_user.role.name != "Doctor":
            raise HTTPException(status_code=403, detail="Only patients and doctors can view summaries")
        
        summaries = db.query(PatientSummary).filter(
            PatientSummary.patient_user_id == patient_user_id
        ).order_by(PatientSummary.summary_date.desc()).all()
        
        return summaries
    
    @staticmethod
    def update_patient_summary(
        db: Session,
        summary_id: int,
        update_data: PatientSummaryUpdate,
        current_user: User
    ):
        """Update patient summary (doctor only)"""
        
        summary = db.query(PatientSummary).filter(
            PatientSummary.id == summary_id
        ).first()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Patient summary not found")
        
        # Only creating doctor can update
        if current_user.role.name != "Doctor" or summary.doctor_user_id != current_user.id:
            raise HTTPException(status_code=403, detail="Only the creating doctor can update this summary")
        
        # Update fields
        if update_data.title is not None:
            summary.title = update_data.title
        if update_data.doctor_notes is not None:
            summary.doctor_notes = update_data.doctor_notes
        if update_data.special_instructions is not None:
            summary.special_instructions = update_data.special_instructions
        
        summary.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(summary)
        
        # Regenerate documents
        PatientSummaryService.generate_documents(db, summary_id)
        
        return summary
