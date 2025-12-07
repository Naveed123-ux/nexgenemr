from sqlalchemy.orm import Session, joinedload
from models.discharge_summary_model import DischargeSummary
from models.user_model import User
from models.patient_profile_model import PatientProfile
from models.appointment_model import Appointment
from models.soap_note_model import SoapNote
from models.prescription_model import Prescription
from models.lab_result_model import LabResult
from models.clinical_data_model import Vitals, MedicalHistory
from models.appointment_icd_code_model import AppointmentICDCode
from models.icd_code_model import ICDCode
from models.hospital_model import Hospital
from schemas.discharge_summary_schema import DischargeSummaryCreate, DischargeSummaryUpdate
from fastapi import HTTPException
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import os


class DischargeSummaryService:
    
    @staticmethod
    def get_patient_comprehensive_data(db: Session, patient_user_id: int):
        """Gather all patient information for discharge summary"""
        
        # Get patient user and profile with careful relationship loading
        patient_user = db.query(User).filter(User.id == patient_user_id).first()
        
        if not patient_user:
            raise HTTPException(status_code=404, detail="Patient user not found")
        
        # Get patient profile separately to ensure relationship exists
        patient_profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == patient_user_id
        ).first()
        
        if not patient_profile:
            raise HTTPException(status_code=404, detail="Patient profile not found")
        
        # Get all appointments with safe relationship loading
        appointments = db.query(Appointment).filter(
            Appointment.patient_profile_id == patient_profile.id
        ).order_by(Appointment.id.desc()).all()
        
        # Load related data for each appointment
        for apt in appointments:
            # Load doctor if relationship exists
            if apt.doctor_user_id:
                apt.doctor = db.query(User).filter(User.id == apt.doctor_user_id).first()
            
            # Load SOAP note if exists
            apt.soap_note = db.query(SoapNote).filter(
                SoapNote.appointment_id == apt.id
            ).first()
            
            # Load vitals if exists
            apt.vitals = db.query(Vitals).filter(
                Vitals.appointment_id == apt.id
            ).first()
            
            # Load ICD code if exists
            if apt.icd_code_id:
                apt.icd_code = db.query(ICDCode).filter(ICDCode.id == apt.icd_code_id).first()
            
            # Load appointment_icd_codes junction table entries
            apt.appointment_icd_codes = db.query(AppointmentICDCode).filter(
                AppointmentICDCode.appointment_id == apt.id
            ).all()
            
            # Load ICD codes for each junction entry
            for apt_icd in apt.appointment_icd_codes:
                if apt_icd.icd_code_id:
                    apt_icd.icd_code = db.query(ICDCode).filter(ICDCode.id == apt_icd.icd_code_id).first()
        
        # Get prescriptions
        prescriptions = db.query(Prescription).filter(
            Prescription.patient_user_id == patient_user_id
        ).all()
        
        # Get lab results
        lab_results = db.query(LabResult).filter(
            LabResult.patient_user_id == patient_user_id
        ).all()
        
        # Get medical history
        medical_history = db.query(MedicalHistory).filter(
            MedicalHistory.patient_profile_id == patient_profile.id
        ).first()
        
        # Get latest vitals from appointments
        latest_vitals = None
        for apt in appointments:
            if apt.vitals:
                latest_vitals = apt.vitals
                break
        
        return {
            "patient_user": patient_user,
            "patient_profile": patient_profile,
            "appointments": appointments,
            "prescriptions": prescriptions,
            "lab_results": lab_results,
            "medical_history": medical_history,
            "latest_vitals": latest_vitals
        }
    
    @staticmethod
    def auto_generate_summary_content(patient_data: dict) -> dict:
        """Auto-generate all discharge summary content from patient data"""
        
        patient_profile = patient_data["patient_profile"]
        appointments = patient_data["appointments"]
        prescriptions = patient_data["prescriptions"]
        lab_results = patient_data["lab_results"]
        medical_history = patient_data["medical_history"]
        latest_vitals = patient_data["latest_vitals"]
        
        # Auto-generate chief complaint from patient profile or first appointment
        chief_complaint = patient_profile.chief_complaint if patient_profile.chief_complaint else "Not documented"
        
        # Auto-generate history from SOAP notes
        history_parts = []
        for apt in appointments:
            if apt.soap_note and apt.soap_note.subjective:
                history_parts.append(f"- {apt.soap_note.subjective}")
        history_of_present_illness = "\n".join(history_parts) if history_parts else "No documented history"
        
        # Auto-generate past medical history
        past_medical_history = "Not documented"
        if medical_history and medical_history.past_medical_history:
            if isinstance(medical_history.past_medical_history, list):
                past_medical_history = ", ".join(medical_history.past_medical_history)
            else:
                past_medical_history = str(medical_history.past_medical_history)
        
        # Auto-generate medications on admission
        medications_on_admission = "None documented"
        if medical_history and medical_history.current_medications:
            if isinstance(medical_history.current_medications, list):
                medications_on_admission = "\n".join([f"- {med}" for med in medical_history.current_medications])
            else:
                medications_on_admission = str(medical_history.current_medications)
        
        # Auto-generate allergies
        allergies = "No known allergies"
        if medical_history and medical_history.allergies:
            if isinstance(medical_history.allergies, list):
                allergies = ", ".join(medical_history.allergies)
            else:
                allergies = str(medical_history.allergies)
        
        # Auto-generate physical examination from SOAP notes
        exam_parts = []
        for apt in appointments:
            if apt.soap_note and apt.soap_note.objective:
                exam_parts.append(f"- {apt.soap_note.objective}")
        physical_examination = "\n".join(exam_parts) if exam_parts else "No documented examination"
        
        # Auto-generate vital signs
        vital_signs = "Not documented"
        if latest_vitals:
            vital_parts = []
            if latest_vitals.blood_pressure:
                vital_parts.append(f"BP: {latest_vitals.blood_pressure}")
            if latest_vitals.heart_rate:
                vital_parts.append(f"HR: {latest_vitals.heart_rate} bpm")
            if latest_vitals.respiratory_rate:
                vital_parts.append(f"RR: {latest_vitals.respiratory_rate} /min")
            if latest_vitals.temperature:
                vital_parts.append(f"Temp: {latest_vitals.temperature}°F")
            if latest_vitals.oxygen_saturation:
                vital_parts.append(f"O2 Sat: {latest_vitals.oxygen_saturation}%")
            vital_signs = ", ".join(vital_parts) if vital_parts else "Not documented"
        
        # Auto-generate hospital course from appointments
        course_parts = []
        for apt in appointments:
            if apt.soap_note and apt.soap_note.assessment:
                course_parts.append(f"- {apt.soap_note.assessment}")
        hospital_course = "\n".join(course_parts) if course_parts else "No documented course"
        
        # Auto-generate procedures from appointment reasons
        procedures = []
        for apt in appointments:
            if apt.reason_for_visit:
                procedures.append(f"- {apt.reason_for_visit}")
        procedures_performed = "\n".join(procedures) if procedures else "No procedures documented"
        
        # Auto-generate lab results summary
        lab_summary_parts = []
        for lab in lab_results:
            # LabResult has specific fields: hemoglobin, wbc, glucose, creatinine, alt, total_cholesterol
            if lab.hemoglobin:
                lab_summary_parts.append(f"- Hemoglobin: {lab.hemoglobin} g/dL")
            if lab.wbc:
                lab_summary_parts.append(f"- WBC: {lab.wbc} x10^9/L")
            if lab.glucose:
                lab_summary_parts.append(f"- Glucose: {lab.glucose} mg/dL")
            if lab.creatinine:
                lab_summary_parts.append(f"- Creatinine: {lab.creatinine} mg/dL")
            if lab.alt:
                lab_summary_parts.append(f"- ALT: {lab.alt} U/L")
            if lab.total_cholesterol:
                lab_summary_parts.append(f"- Total Cholesterol: {lab.total_cholesterol} mg/dL")
            if lab.result_date:
                lab_summary_parts.append(f"  (Date: {lab.result_date})")
        lab_results_summary = "\n".join(lab_summary_parts) if lab_summary_parts else "No lab results available"
        
        # Auto-generate diagnoses from ICD codes
        diagnoses = []
        for apt in appointments:
            if apt.icd_code:
                diagnoses.append(f"{apt.icd_code.code} - {apt.icd_code.description}")
            # Check appointment_icd_codes relationship
            if hasattr(apt, 'appointment_icd_codes'):
                for apt_icd in apt.appointment_icd_codes:
                    if hasattr(apt_icd, 'icd_code') and apt_icd.icd_code:
                        diagnoses.append(f"{apt_icd.icd_code.code} - {apt_icd.icd_code.description}")
        
        primary_diagnosis = diagnoses[0] if diagnoses else "No diagnosis documented"
        secondary_diagnosis = ", ".join(diagnoses[1:]) if len(diagnoses) > 1 else "None"
        
        # Auto-generate condition on discharge
        condition_on_discharge = "Stable" if appointments else "Not documented"
        
        # Auto-generate discharge medications from prescriptions
        discharge_meds = []
        for rx in prescriptions:
            if rx.status == "active" or rx.status == "ACTIVE":
                discharge_meds.append(f"- {rx.medication} {rx.dosage} {rx.frequency}")
        discharge_medications = "\n".join(discharge_meds) if discharge_meds else "No medications prescribed"
        
        # Auto-generate discharge instructions from SOAP plans
        instructions = []
        for apt in appointments:
            if apt.soap_note and apt.soap_note.plan:
                instructions.append(f"- {apt.soap_note.plan}")
        discharge_instructions = "\n".join(instructions) if instructions else "Follow up with primary care physician"
        
        # Auto-generate follow-up instructions
        follow_up_instructions = "Follow up with primary care physician in 1-2 weeks or as directed"
        
        # Auto-generate diet instructions
        diet_instructions = "Regular diet as tolerated"
        
        # Auto-generate activity restrictions
        activity_restrictions = "Resume normal activities as tolerated"
        
        # Complications
        complications = "None documented"
        
        # Consultations
        consultations = "None documented"
        
        # Additional notes
        additional_notes = f"Total appointments during admission: {len(appointments)}"
        
        return {
            "chief_complaint": chief_complaint,
            "history_of_present_illness": history_of_present_illness,
            "past_medical_history": past_medical_history,
            "medications_on_admission": medications_on_admission,
            "allergies": allergies,
            "physical_examination": physical_examination,
            "vital_signs": vital_signs,
            "hospital_course": hospital_course,
            "procedures_performed": procedures_performed,
            "lab_results_summary": lab_results_summary,
            "primary_diagnosis": primary_diagnosis,
            "secondary_diagnosis": secondary_diagnosis,
            "condition_on_discharge": condition_on_discharge,
            "discharge_medications": discharge_medications,
            "discharge_instructions": discharge_instructions,
            "follow_up_instructions": follow_up_instructions,
            "diet_instructions": diet_instructions,
            "activity_restrictions": activity_restrictions,
            "complications": complications,
            "consultations": consultations,
            "additional_notes": additional_notes
        }
    
    @staticmethod
    def create_discharge_summary(
        db: Session,
        summary_data: DischargeSummaryCreate,
        current_user: User
    ):
        """Create a new discharge summary - AUTO-GENERATED from database"""
        
        # Verify role (Doctor or Hospital_Admin)
        if current_user.role.name not in ["Doctor", "Hospital_Admin"]:
            raise HTTPException(
                status_code=403,
                detail="Only doctors and hospital admins can create discharge summaries"
            )
        
        # Get patient data
        patient_data = DischargeSummaryService.get_patient_comprehensive_data(
            db, summary_data.patient_user_id
        )
        
        patient_user = patient_data["patient_user"]
        patient_profile = patient_data["patient_profile"]
        
        # Get hospital ID based on user role
        hospital_id = None
        if current_user.role.name == "Doctor":
            if current_user.doctor_profile and current_user.doctor_profile.department:
                hospital_id = current_user.doctor_profile.department.hospital_id
            else:
                raise HTTPException(status_code=400, detail="Doctor must be associated with a hospital")
        elif current_user.role.name == "Hospital_Admin":
            # Get hospital for admin
            hospital = db.query(Hospital).filter(
                Hospital.admin_user_id == current_user.id
            ).first()
            if not hospital:
                raise HTTPException(status_code=400, detail="Hospital not found for this admin")
            hospital_id = hospital.id
        
        # Admission date = patient user creation date
        admission_date = datetime.fromisoformat(patient_user.created_at) if isinstance(patient_user.created_at, str) else patient_user.created_at
        
        # Discharge date = today
        discharge_date = datetime.utcnow()
        
        # AUTO-GENERATE all content from database
        auto_content = DischargeSummaryService.auto_generate_summary_content(patient_data)
        
        # Create discharge summary with auto-generated content
        discharge_summary = DischargeSummary(
            patient_user_id=summary_data.patient_user_id,
            hospital_id=hospital_id,
            created_by_doctor_id=current_user.id,
            admission_date=admission_date,
            discharge_date=discharge_date,
            **auto_content  # Unpack all auto-generated fields
        )
        
        db.add(discharge_summary)
        db.commit()
        db.refresh(discharge_summary)
        
        # Generate PDF and Word documents
        DischargeSummaryService.generate_documents(db, discharge_summary.id)
        
        return discharge_summary
    
    @staticmethod
    def generate_documents(db: Session, summary_id: int):
        """Generate PDF and Word documents for discharge summary"""
        
        summary = db.query(DischargeSummary).options(
            joinedload(DischargeSummary.patient),
            joinedload(DischargeSummary.hospital),
            joinedload(DischargeSummary.created_by_doctor)
        ).filter(DischargeSummary.id == summary_id).first()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Discharge summary not found")
        
        # Get patient profile
        patient_profile = db.query(PatientProfile).filter(
            PatientProfile.user_id == summary.patient_user_id
        ).first()
        
        # Create output directory
        output_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'discharge_summaries')
        os.makedirs(output_dir, exist_ok=True)
        
        # Generate filenames
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"discharge_summary_{summary.patient_user_id}_{timestamp}.pdf"
        word_filename = f"discharge_summary_{summary.patient_user_id}_{timestamp}.docx"
        
        pdf_path = os.path.join(output_dir, pdf_filename)
        word_path = os.path.join(output_dir, word_filename)
        
        # Generate PDF
        DischargeSummaryService._generate_pdf(summary, patient_profile, pdf_path)
        
        # Generate Word document
        DischargeSummaryService._generate_word(summary, patient_profile, word_path)
        
        # Update summary with file paths
        summary.pdf_file_path = f"/static/discharge_summaries/{pdf_filename}"
        summary.word_file_path = f"/static/discharge_summaries/{word_filename}"
        db.commit()
        
        return summary
    
    @staticmethod
    def _generate_pdf(summary: DischargeSummary, patient_profile: PatientProfile, output_path: str):
        """Generate PDF discharge summary with professional template"""
        
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=1*inch, bottomMargin=0.75*inch)
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1a5490'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#1a5490'),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )
        
        # Add logo if exists
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'static', 'images', 'logos', 'logocure.png')
        if os.path.exists(logo_path):
            # Maintain aspect ratio by calculating proportional height
            pil_img = PILImage.open(logo_path)
            img_width, img_height = pil_img.size
            aspect_ratio = img_height / img_width
            
            # Set desired width and calculate proportional height
            desired_width = 2.5 * inch
            desired_height = desired_width * aspect_ratio
            
            logo = Image(logo_path, width=desired_width, height=desired_height)
            story.append(logo)
            story.append(Spacer(1, 0.2*inch))
        
        # Header
        story.append(Paragraph("DISCHARGE SUMMARY", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Hospital and patient information table
        patient_info_data = [
            ['Hospital:', summary.hospital.name if summary.hospital else 'N/A'],
            ['Patient Name:', f"{summary.patient.first_name} {summary.patient.last_name}"],
            ['Patient ID:', str(summary.patient_user_id)],
            ['Date of Birth:', patient_profile.subscriber_dob.strftime('%B %d, %Y') if patient_profile.subscriber_dob else 'N/A'],
            ['Admission Date:', summary.admission_date.strftime('%B %d, %Y')],
            ['Discharge Date:', summary.discharge_date.strftime('%B %d, %Y')],
            ['Attending Physician:', f"Dr. {summary.created_by_doctor.first_name} {summary.created_by_doctor.last_name}"]
        ]
        
        patient_info_table = Table(patient_info_data, colWidths=[2*inch, 4.5*inch])
        patient_info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8f4f8')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ]))
        
        story.append(patient_info_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Add sections
        sections = [
            ('CHIEF COMPLAINT', summary.chief_complaint),
            ('HISTORY OF PRESENT ILLNESS', summary.history_of_present_illness),
            ('PAST MEDICAL HISTORY', summary.past_medical_history),
            ('MEDICATIONS ON ADMISSION', summary.medications_on_admission),
            ('ALLERGIES', summary.allergies),
            ('PHYSICAL EXAMINATION', summary.physical_examination),
            ('VITAL SIGNS', summary.vital_signs),
            ('HOSPITAL COURSE', summary.hospital_course),
            ('PROCEDURES PERFORMED', summary.procedures_performed),
            ('LABORATORY RESULTS', summary.lab_results_summary),
            ('PRIMARY DIAGNOSIS', summary.primary_diagnosis),
            ('SECONDARY DIAGNOSIS', summary.secondary_diagnosis),
            ('CONDITION ON DISCHARGE', summary.condition_on_discharge),
            ('DISCHARGE MEDICATIONS', summary.discharge_medications),
            ('DISCHARGE INSTRUCTIONS', summary.discharge_instructions),
            ('FOLLOW-UP INSTRUCTIONS', summary.follow_up_instructions),
            ('DIET INSTRUCTIONS', summary.diet_instructions),
            ('ACTIVITY RESTRICTIONS', summary.activity_restrictions),
            ('COMPLICATIONS', summary.complications),
            ('CONSULTATIONS', summary.consultations),
            ('ADDITIONAL NOTES', summary.additional_notes)
        ]
        
        for section_title, section_content in sections:
            if section_content:
                story.append(Paragraph(section_title, heading_style))
                story.append(Paragraph(section_content or 'N/A', normal_style))
                story.append(Spacer(1, 0.15*inch))
        
        # Footer with signature
        story.append(Spacer(1, 0.5*inch))
        signature_data = [
            ['', ''],
            ['_' * 40, '_' * 40],
            [f"Dr. {summary.created_by_doctor.first_name} {summary.created_by_doctor.last_name}", 'Date: ' + summary.discharge_date.strftime('%B %d, %Y')],
            ['Attending Physician', '']
        ]
        
        signature_table = Table(signature_data, colWidths=[3*inch, 3*inch])
        signature_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 2), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 2), (-1, -1), 5)
        ]))
        
        story.append(signature_table)
        
        # Add e-signatures section if any signatures exist
        from models.signature_model import UserSignature
        from sqlalchemy.orm import Session
        
        if summary.signed_by_doctor_id or summary.signed_by_staff_id or summary.signed_by_admin_id:
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph("Electronic Signatures:", heading_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Doctor signature
            if summary.signed_by_doctor_id and summary.signed_by_doctor:
                doctor_sig = None
                try:
                    from db.db import SessionLocal
                    db = SessionLocal()
                    doctor_sig = db.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_doctor_id,
                        UserSignature.is_active == True
                    ).first()
                    db.close()
                except:
                    pass
                
                if doctor_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        doctor_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            sig_img = Image(sig_path, width=2*inch, height=0.75*inch)
                            story.append(sig_img)
                        except:
                            pass
                        
                        doctor_name = f"Dr. {summary.signed_by_doctor.first_name} {summary.signed_by_doctor.last_name}"
                        story.append(Paragraph(f"<b>{doctor_name}</b>", normal_style))
                        if summary.doctor_signed_at:
                            story.append(Paragraph(
                                f"Signed: {summary.doctor_signed_at.strftime('%B %d, %Y at %I:%M %p')}", 
                                normal_style
                            ))
                        story.append(Spacer(1, 0.3*inch))
            
            # Staff signature
            if summary.signed_by_staff_id and summary.signed_by_staff:
                staff_sig = None
                try:
                    from db.db import SessionLocal
                    db = SessionLocal()
                    staff_sig = db.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_staff_id,
                        UserSignature.is_active == True
                    ).first()
                    db.close()
                except:
                    pass
                
                if staff_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        staff_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            sig_img = Image(sig_path, width=2*inch, height=0.75*inch)
                            story.append(sig_img)
                        except:
                            pass
                        
                        staff_name = f"{summary.signed_by_staff.first_name} {summary.signed_by_staff.last_name}"
                        story.append(Paragraph(f"<b>{staff_name} (Staff)</b>", normal_style))
                        if summary.staff_signed_at:
                            story.append(Paragraph(
                                f"Signed: {summary.staff_signed_at.strftime('%B %d, %Y at %I:%M %p')}", 
                                normal_style
                            ))
                        story.append(Spacer(1, 0.3*inch))
            
            # Admin signature
            if summary.signed_by_admin_id and summary.signed_by_admin:
                admin_sig = None
                try:
                    from db.db import SessionLocal
                    db = SessionLocal()
                    admin_sig = db.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_admin_id,
                        UserSignature.is_active == True
                    ).first()
                    db.close()
                except:
                    pass
                
                if admin_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        admin_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            sig_img = Image(sig_path, width=2*inch, height=0.75*inch)
                            story.append(sig_img)
                        except:
                            pass
                        
                        admin_name = f"{summary.signed_by_admin.first_name} {summary.signed_by_admin.last_name}"
                        story.append(Paragraph(f"<b>{admin_name} (Hospital Admin)</b>", normal_style))
                        if summary.admin_signed_at:
                            story.append(Paragraph(
                                f"Signed: {summary.admin_signed_at.strftime('%B %d, %Y at %I:%M %p')}", 
                                normal_style
                            ))
                        story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(story)
    
    @staticmethod
    def _generate_word(summary: DischargeSummary, patient_profile: PatientProfile, output_path: str):
        """Generate Word discharge summary with professional template"""
        from docx.shared import Inches
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add logo if exists
        logo_path = os.path.join(os.path.dirname(__file__), '..', 'static', 'images', 'logos', 'logocure.png')
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(2.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title = doc.add_heading('DISCHARGE SUMMARY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(26, 84, 144)
        
        doc.add_paragraph()
        
        # Patient Information Table
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Hospital:', summary.hospital.name if summary.hospital else 'N/A'),
            ('Patient Name:', f"{summary.patient.first_name} {summary.patient.last_name}"),
            ('Patient ID:', str(summary.patient_user_id)),
            ('Date of Birth:', patient_profile.subscriber_dob.strftime('%B %d, %Y') if patient_profile.subscriber_dob else 'N/A'),
            ('Admission Date:', summary.admission_date.strftime('%B %d, %Y')),
            ('Discharge Date:', summary.discharge_date.strftime('%B %d, %Y')),
            ('Attending Physician:', f"Dr. {summary.created_by_doctor.first_name} {summary.created_by_doctor.last_name}")
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Add sections
        sections_data = [
            ('CHIEF COMPLAINT', summary.chief_complaint),
            ('HISTORY OF PRESENT ILLNESS', summary.history_of_present_illness),
            ('PAST MEDICAL HISTORY', summary.past_medical_history),
            ('MEDICATIONS ON ADMISSION', summary.medications_on_admission),
            ('ALLERGIES', summary.allergies),
            ('PHYSICAL EXAMINATION', summary.physical_examination),
            ('VITAL SIGNS', summary.vital_signs),
            ('HOSPITAL COURSE', summary.hospital_course),
            ('PROCEDURES PERFORMED', summary.procedures_performed),
            ('LABORATORY RESULTS', summary.lab_results_summary),
            ('PRIMARY DIAGNOSIS', summary.primary_diagnosis),
            ('SECONDARY DIAGNOSIS', summary.secondary_diagnosis),
            ('CONDITION ON DISCHARGE', summary.condition_on_discharge),
            ('DISCHARGE MEDICATIONS', summary.discharge_medications),
            ('DISCHARGE INSTRUCTIONS', summary.discharge_instructions),
            ('FOLLOW-UP INSTRUCTIONS', summary.follow_up_instructions),
            ('DIET INSTRUCTIONS', summary.diet_instructions),
            ('ACTIVITY RESTRICTIONS', summary.activity_restrictions),
            ('COMPLICATIONS', summary.complications),
            ('CONSULTATIONS', summary.consultations),
            ('ADDITIONAL NOTES', summary.additional_notes)
        ]
        
        for section_title, section_content in sections_data:
            if section_content:
                heading = doc.add_heading(section_title, level=2)
                heading_run = heading.runs[0]
                heading_run.font.color.rgb = RGBColor(26, 84, 144)
                
                doc.add_paragraph(section_content or 'N/A')
        
        # Signature section
        doc.add_paragraph()
        doc.add_paragraph()
        
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.rows[0].cells[0].text = '_' * 40
        sig_table.rows[0].cells[1].text = '_' * 40
        sig_table.rows[1].cells[0].text = f"Dr. {summary.created_by_doctor.first_name} {summary.created_by_doctor.last_name}\nAttending Physician"
        sig_table.rows[1].cells[1].text = f"Date: {summary.discharge_date.strftime('%B %d, %Y')}"
        
        for row in sig_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add e-signatures section if any signatures exist
        from models.signature_model import UserSignature
        
        if summary.signed_by_doctor_id or summary.signed_by_staff_id or summary.signed_by_admin_id:
            doc.add_paragraph()
            doc.add_heading('Electronic Signatures', level=2)
            
            # Doctor signature
            if summary.signed_by_doctor_id and summary.signed_by_doctor:
                doctor_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    doctor_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_doctor_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if doctor_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        doctor_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            doc.add_picture(sig_path, width=Inches(2))
                        except:
                            pass
                        
                        doctor_name = f"Dr. {summary.signed_by_doctor.first_name} {summary.signed_by_doctor.last_name}"
                        p = doc.add_paragraph()
                        p.add_run(doctor_name).bold = True
                        if summary.doctor_signed_at:
                            doc.add_paragraph(f"Signed: {summary.doctor_signed_at.strftime('%B %d, %Y at %I:%M %p')}")
                        doc.add_paragraph()
            
            # Staff signature
            if summary.signed_by_staff_id and summary.signed_by_staff:
                staff_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    staff_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_staff_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if staff_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        staff_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            doc.add_picture(sig_path, width=Inches(2))
                        except:
                            pass
                        
                        staff_name = f"{summary.signed_by_staff.first_name} {summary.signed_by_staff.last_name}"
                        p = doc.add_paragraph()
                        p.add_run(f"{staff_name} (Staff)").bold = True
                        if summary.staff_signed_at:
                            doc.add_paragraph(f"Signed: {summary.staff_signed_at.strftime('%B %d, %Y at %I:%M %p')}")
                        doc.add_paragraph()
            
            # Admin signature
            if summary.signed_by_admin_id and summary.signed_by_admin:
                admin_sig = None
                try:
                    from db.db import SessionLocal
                    db_session = SessionLocal()
                    admin_sig = db_session.query(UserSignature).filter(
                        UserSignature.user_id == summary.signed_by_admin_id,
                        UserSignature.is_active == True
                    ).first()
                    db_session.close()
                except:
                    pass
                
                if admin_sig:
                    sig_path = os.path.join(
                        os.path.dirname(__file__), 
                        '..', 
                        admin_sig.signature_file_path.lstrip('/')
                    )
                    if os.path.exists(sig_path):
                        try:
                            doc.add_picture(sig_path, width=Inches(2))
                        except:
                            pass
                        
                        admin_name = f"{summary.signed_by_admin.first_name} {summary.signed_by_admin.last_name}"
                        p = doc.add_paragraph()
                        p.add_run(f"{admin_name} (Hospital Admin)").bold = True
                        if summary.admin_signed_at:
                            doc.add_paragraph(f"Signed: {summary.admin_signed_at.strftime('%B %d, %Y at %I:%M %p')}")
                        doc.add_paragraph()
        
        # Save document
        doc.save(output_path)
    
    @staticmethod
    def get_discharge_summary(db: Session, summary_id: int, current_user: User):
        """Get a discharge summary by ID"""
        
        summary = db.query(DischargeSummary).options(
            joinedload(DischargeSummary.patient),
            joinedload(DischargeSummary.hospital),
            joinedload(DischargeSummary.created_by_doctor)
        ).filter(DischargeSummary.id == summary_id).first()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Discharge summary not found")
        
        # Check permissions
        if current_user.role.name == "Patient":
            if summary.patient_user_id != current_user.id:
                raise HTTPException(status_code=403, detail="You can only view your own discharge summaries")
        elif current_user.role.name == "Doctor":
            # Doctors can view summaries from their hospital
            if current_user.doctor_profile.department.hospital_id != summary.hospital_id:
                raise HTTPException(status_code=403, detail="You can only view discharge summaries from your hospital")
        elif current_user.role.name == "Hospital_Admin":
            # Use explicit query to avoid lazy-loading issues
            hospital = db.query(Hospital).filter(
                Hospital.admin_user_id == current_user.id
            ).first()
            if not hospital or hospital.id != summary.hospital_id:
                raise HTTPException(status_code=403, detail="You can only view discharge summaries from your hospital")
        else:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        
        return summary
    
    @staticmethod
    def get_patient_discharge_summaries(db: Session, patient_user_id: int, current_user: User):
        """Get all discharge summaries for a patient"""
        
        # Check permissions
        if current_user.role.name == "Patient":
            if patient_user_id != current_user.id:
                raise HTTPException(status_code=403, detail="You can only view your own discharge summaries")
        
        summaries = db.query(DischargeSummary).options(
            joinedload(DischargeSummary.hospital),
            joinedload(DischargeSummary.created_by_doctor)
        ).filter(
            DischargeSummary.patient_user_id == patient_user_id
        ).order_by(DischargeSummary.discharge_date.desc()).all()
        
        return summaries
    
    @staticmethod
    def update_discharge_summary(
        db: Session,
        summary_id: int,
        update_data: DischargeSummaryUpdate,
        current_user: User
    ):
        """Update a discharge summary (only if not finalized)"""
        
        summary = db.query(DischargeSummary).filter(
            DischargeSummary.id == summary_id
        ).first()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Discharge summary not found")
        
        # Only the creating doctor can update
        if summary.created_by_doctor_id != current_user.id:
            raise HTTPException(status_code=403, detail="Only the creating doctor can update this summary")
        
        # Cannot update if finalized
        if summary.is_finalized:
            raise HTTPException(status_code=400, detail="Cannot update a finalized discharge summary")
        
        # Update fields
        update_dict = update_data.dict(exclude_unset=True)
        for key, value in update_dict.items():
            setattr(summary, key, value)
        
        summary.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(summary)
        
        # Regenerate documents
        DischargeSummaryService.generate_documents(db, summary_id)
        
        return summary
    
    @staticmethod
    def finalize_discharge_summary(db: Session, summary_id: int, current_user: User):
        """Finalize a discharge summary (makes it read-only)"""
        
        summary = db.query(DischargeSummary).filter(
            DischargeSummary.id == summary_id
        ).first()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Discharge summary not found")
        
        # Only the creating doctor can finalize
        if summary.created_by_doctor_id != current_user.id:
            raise HTTPException(status_code=403, detail="Only the creating doctor can finalize this summary")
        
        summary.is_finalized = True
        summary.updated_at = datetime.utcnow()
        db.commit()
        
        return summary
    
    @staticmethod
    def delete_discharge_summary(db: Session, summary_id: int, current_user: User):
        """Delete a discharge summary (only if not finalized)"""
        
        summary = db.query(DischargeSummary).filter(
            DischargeSummary.id == summary_id
        ).first()
        
        if not summary:
            raise HTTPException(status_code=404, detail="Discharge summary not found")
        
        # Only the creating doctor can delete
        if summary.created_by_doctor_id != current_user.id:
            raise HTTPException(status_code=403, detail="Only the creating doctor can delete this summary")
        
        # Cannot delete if finalized
        if summary.is_finalized:
            raise HTTPException(status_code=400, detail="Cannot delete a finalized discharge summary")
        
        # Delete associated files
        if summary.pdf_file_path:
            pdf_full_path = os.path.join(os.path.dirname(__file__), '..', summary.pdf_file_path.lstrip('/'))
            if os.path.exists(pdf_full_path):
                os.remove(pdf_full_path)
        
        if summary.word_file_path:
            word_full_path = os.path.join(os.path.dirname(__file__), '..', summary.word_file_path.lstrip('/'))
            if os.path.exists(word_full_path):
                os.remove(word_full_path)
        
        db.delete(summary)
        db.commit()
        
        return {"message": "Discharge summary deleted successfully"}
